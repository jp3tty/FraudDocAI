"""
FraudDocAI - AI Service
FastAPI service for document processing and fraud detection
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer
import uvicorn
import logging
from typing import List, Dict, Any
import asyncio
from datetime import datetime
from PIL import Image, ImageEnhance, ImageFilter
import io
from contextlib import asynccontextmanager
from config import config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variables for AI models
document_processor = None
fraud_detector = None
embedding_model = None
text_classifier = None
question_answering_model = None
document_qa_service = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize AI models on startup and cleanup on shutdown"""
    global document_processor, fraud_detector, embedding_model, text_classifier, question_answering_model, document_qa_service
    
    logger.info("Starting FraudDocAI AI Service...")
    
    # Get AI model configuration
    ai_config = config.get_ai_config()
    
    try:
        # Initialize models with real Hugging Face models
        logger.info("Loading AI models...")
        
        # Import transformers
        from transformers import (
            pipeline, 
            AutoTokenizer, 
            AutoModelForSequenceClassification,
            AutoModelForQuestionAnswering,
            AutoModel
        )
        from sentence_transformers import SentenceTransformer
        
        # 1. Text Classification for fraud detection
        logger.info(f"Loading fraud detection classifier: {ai_config['emotion_model']}")
        text_classifier = pipeline(
            "text-classification",
            model=ai_config['emotion_model'],
            return_all_scores=True
        )
        
        # 2. Document Question Answering
        logger.info(f"Loading question answering model: {ai_config['qa_model']}")
        question_answering_model = pipeline(
            "question-answering",
            model=ai_config['qa_model']
        )
        
        # 3. Sentence Embeddings
        logger.info(f"Loading sentence transformer: {ai_config['embedding_model']}")
        embedding_model = SentenceTransformer(ai_config['embedding_model'])
        
        # 4. Document processor (OCR + text extraction)
        logger.info("Initializing document processor...")
        document_processor = "ready"  # We'll implement OCR functionality
        
        # 5. Fraud detector (emotion + pattern analysis)
        logger.info("Initializing fraud detector...")
        fraud_detector = "ready"  # We'll implement fraud detection
        
        # 6. Document QA Service
        logger.info("Initializing document QA service...")
        from document_qa_service import DocumentQuestionAnsweringService
        document_qa_service = DocumentQuestionAnsweringService()
        
        logger.info("✅ All AI models loaded successfully!")
        logger.info(f"Models loaded: {len([m for m in [text_classifier, question_answering_model, embedding_model, document_qa_service] if m is not None])}/4")
        
    except Exception as e:
        logger.error(f"Failed to start AI Service: {e}")
        # Don't raise - start with basic functionality
        logger.warning("Starting with limited AI functionality...")
        document_processor = "limited"
        fraud_detector = "limited"
        embedding_model = "limited"
        text_classifier = "limited"
        question_answering_model = "limited"
        document_qa_service = "limited"
    
    yield  # This is where the app runs
    
    # Cleanup on shutdown
    logger.info("Shutting down AI Service...")

# Initialize FastAPI app with lifespan
app = FastAPI(
    title="FraudDocAI AI Service",
    description="AI-powered document fraud detection service",
    version="1.0.0",
    lifespan=lifespan
)

# CORS middleware with configuration
cors_config = config.get_cors_config()
app.add_middleware(
    CORSMiddleware,
    **cors_config
)

# Security
security = HTTPBearer()

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "service": "FraudDocAI AI Service",
        "status": "running",
        "version": "1.0.0",
        "config": {
            "server": config.get_server_config(),
            "ai_models": config.get_ai_config()
        }
    }

@app.get("/health")
async def health_check():
    """Detailed health check endpoint"""
    global document_processor, fraud_detector, embedding_model, text_classifier, question_answering_model, document_qa_service
    
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "services": {
            "document_processor": "ready" if document_processor == "ready" else "limited",
            "fraud_detector": "ready" if fraud_detector == "ready" else "limited",
            "embedding_model": "ready" if embedding_model and embedding_model != "limited" else "limited",
            "text_classifier": "ready" if text_classifier and text_classifier != "limited" else "limited",
            "question_answering": "ready" if question_answering_model and question_answering_model != "limited" else "limited",
            "document_qa_service": "ready" if document_qa_service and document_qa_service != "limited" else "limited"
        },
        "configuration": config.get_all_config()
    }

@app.post("/process-document")
async def process_document(
    file: UploadFile = File(...),
    token: str = Depends(security)
):
    """
    Process uploaded document for fraud detection
    """
    try:
        logger.info(f"Processing document: {file.filename}")
        
        # Validate file type
        allowed_types = [
            "application/pdf",
            "image/jpeg",
            "image/png",
            "image/tiff",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]
        
        if file.content_type not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file.content_type}"
            )
        
        # Read file content
        content = await file.read()
        
        # Extract text from document with quality enhancement
        ocr_result = await extract_text_with_quality_enhancement(content, file.content_type)
        extracted_text = ocr_result["extracted_text"]
        
        # Analyze text for fraud patterns using AI models
        fraud_analysis = await analyze_text_for_fraud(extracted_text)
        
        result = {
            "document_id": f"doc-{datetime.utcnow().strftime('%Y%m%d%H%M%S')}",
            "filename": file.filename,
            "file_type": file.content_type,
            "file_size": len(content),
            "extracted_text": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
            "fraud_score": fraud_analysis["fraud_score"],
            "fraud_risk_level": fraud_analysis["risk_level"],
            "detected_patterns": fraud_analysis["patterns"],
            "processing_time_ms": fraud_analysis["processing_time"],
            "timestamp": datetime.utcnow().isoformat(),
            # OCR Quality Data
            "ocr_quality": {
                "confidence_score": ocr_result["confidence_score"],
                "quality_level": ocr_result["quality_level"],
                "preprocessing_applied": ocr_result["preprocessing_applied"],
                "text_blocks": ocr_result["text_blocks"],
                "processing_notes": ocr_result["processing_notes"],
                "file_type": ocr_result["file_type"]
            }
        }
        
        logger.info(f"Document processed successfully: {file.filename}")
        return result
        
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-text")
async def analyze_text(
    text: str,
    token: str = Depends(security)
):
    """
    Analyze text for fraud patterns using AI models
    """
    try:
        logger.info(f"Analyzing text: {len(text)} characters")
        
        # Analyze text for fraud patterns
        fraud_analysis = await analyze_text_for_fraud(text)
        
        result = {
            "text_length": len(text),
            "fraud_score": fraud_analysis["fraud_score"],
            "risk_level": fraud_analysis["risk_level"],
            "patterns": fraud_analysis["patterns"],
            "emotion_analysis": fraud_analysis.get("emotion_analysis", {}),
            "pattern_analysis": fraud_analysis.get("pattern_analysis", {}),
            "processing_time_ms": fraud_analysis["processing_time"],
            "timestamp": datetime.utcnow().isoformat()
        }
        
        logger.info(f"Text analysis completed: {fraud_analysis['risk_level']} risk")
        return result
        
    except Exception as e:
        logger.error(f"Error analyzing text: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-embeddings")
async def generate_embeddings(
    text: str,
    token: str = Depends(security)
):
    """
    Generate embeddings for text using sentence transformer
    """
    try:
        logger.info(f"Generating embeddings for text: {len(text)} characters")
        
        if embedding_model == "limited" or embedding_model is None:
            raise HTTPException(
                status_code=503,
                detail="Embedding model not available"
            )
        
        start_time = datetime.utcnow()
        
        try:
            # Generate embeddings
            embeddings = embedding_model.encode(text)
            generation_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            result = {
                "text": text[:100] + "..." if len(text) > 100 else text,
                "embeddings": embeddings.tolist(),
                "embedding_dimension": len(embeddings),
                "generation_time_ms": round(generation_time, 2),
                "timestamp": datetime.utcnow().isoformat()
            }
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            # Fallback to mock embeddings
            embeddings = []
            generation_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            result = {
                "text": text[:100] + "..." if len(text) > 100 else text,
                "embeddings": [],
                "embedding_dimension": 0,
                "generation_time_ms": round(generation_time, 2),
                "error": "Embedding generation failed",
                "timestamp": datetime.utcnow().isoformat()
            }
        
        logger.info(f"Embeddings generated: {len(embeddings)} dimensions")
        return result
        
    except Exception as e:
        logger.error(f"Error generating embeddings: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/fraud-patterns")
async def get_fraud_patterns(token: str = Depends(security)):
    """
    Get available fraud detection patterns
    """
    try:
        patterns = [
            {
                "pattern": "urgency_indicators",
                "keywords": ["urgent", "immediate", "asap", "emergency", "critical"],
                "description": "Detects urgent language that may indicate pressure tactics"
            },
            {
                "pattern": "confidentiality_claims",
                "keywords": ["confidential", "secret", "do not share", "private"],
                "description": "Identifies claims of confidentiality that may be manipulative"
            },
            {
                "pattern": "payment_methods",
                "keywords": ["wire transfer", "offshore", "bitcoin", "gift cards"],
                "description": "Detects suspicious payment methods commonly used in fraud"
            },
            {
                "pattern": "amount_tampering",
                "keywords": ["amount", "total", "sum", "cost", "price"],
                "description": "Identifies potential manipulation of monetary amounts"
            }
        ]
        
        return {
            "patterns": patterns,
            "total_patterns": len(patterns),
            "timestamp": datetime.utcnow().isoformat()
        }
        
    except Exception as e:
        logger.error(f"Error getting fraud patterns: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/document-qa")
async def document_qa(
    question: str = Form(...),
    document_text: str = Form(...),
    token: str = Depends(security)
):
    """
    Answer questions about a document using AI
    """
    try:
        logger.info(f"Processing document Q&A: {len(question)} char question, {len(document_text)} char document")
        
        if document_qa_service == "limited" or document_qa_service is None:
            raise HTTPException(
                status_code=503,
                detail="Document QA service not available"
            )
        
        # Process the question
        answer = document_qa_service.answer_question(question, document_text)
        
        result = {
            "question": question,
            "answer": answer["answer"],
            "confidence": answer["confidence"],
            "context_used": answer["context_used"],
            "processing_time_ms": answer["processing_time_ms"],
            "timestamp": datetime.utcnow().isoformat()
        }
        
        logger.info(f"Document Q&A completed: {answer['confidence']:.1%} confidence")
        return result
        
    except Exception as e:
        logger.error(f"Error processing document question: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/analyze-document-fraud")
async def analyze_document_fraud(
    document_text: str = Form(...),
    token: str = Depends(security)
):
    """
    Comprehensive fraud analysis of a document
    """
    try:
        logger.info(f"Analyzing document for fraud: {len(document_text)} characters")
        
        if document_qa_service == "limited" or document_qa_service is None:
            raise HTTPException(
                status_code=503,
                detail="Document analysis service not available"
            )
        
        # Analyze the document
        analysis = document_qa_service.analyze_document_fraud(document_text)
        
        result = {
            "document_length": len(document_text),
            "fraud_risk_score": analysis["fraud_risk_score"],
            "risk_level": analysis["risk_level"],
            "key_indicators": analysis["key_indicators"],
            "recommendations": analysis["recommendations"],
            "processing_time_ms": analysis["processing_time_ms"],
            "timestamp": datetime.utcnow().isoformat()
        }
        
        logger.info(f"Document fraud analysis completed: {analysis['risk_level']} risk")
        return result
        
    except Exception as e:
        logger.error(f"Error analyzing document for fraud: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/qa-model-info")
async def get_qa_model_info(token: str = Depends(security)):
    """
    Get information about the QA model
    """
    try:
        if document_qa_service == "limited" or document_qa_service is None:
            return {
                "model_loaded": False,
                "model_name": "Not available",
                "status": "limited"
            }
        
        model_info = document_qa_service.get_model_info()
        
        return {
            "model_loaded": True,
            "model_name": model_info["model_name"],
            "model_type": model_info["model_type"],
            "max_length": model_info["max_length"],
            "status": "ready"
        }
        
    except Exception as e:
        logger.error(f"Error getting QA model info: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Helper functions for AI processing
async def extract_text_from_document(content: bytes, content_type: str) -> str:
    """Extract text from various document types using OCR and parsing"""
    try:
        if content_type == "application/pdf":
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        
        elif content_type in ["image/jpeg", "image/png", "image/tiff"]:
            import pytesseract
            from PIL import Image
            import io
            image = Image.open(io.BytesIO(content))
            text = pytesseract.image_to_string(image)
            return text.strip()
        
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        
        else:
            return "Unsupported file type for text extraction"
            
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        return "Error extracting text from document"

async def extract_text_with_quality_enhancement(content: bytes, content_type: str) -> dict:
    """Extract text with quality enhancements and confidence scoring"""
    try:
        if content_type == "application/pdf":
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            page_count = len(pdf_reader.pages)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return {
                "extracted_text": text.strip(),
                "confidence_score": 95.0,  # PDF text extraction is typically very reliable
                "quality_level": "excellent",
                "preprocessing_applied": False,
                "text_blocks": page_count,
                "processing_notes": "PDF text extraction - high reliability",
                "file_type": "pdf"
            }
        
        elif content_type in ["image/jpeg", "image/png", "image/tiff"]:
            import pytesseract
            from PIL import Image, ImageEnhance, ImageFilter
            import io
            
            # Preprocess image for better OCR
            processed_image = preprocess_image_for_ocr(content)
            
            # Extract text with confidence data
            ocr_data = pytesseract.image_to_data(
                processed_image, 
                output_type=pytesseract.Output.DICT,
                config='--psm 6'  # Assume single text block
            )
            
            # Calculate confidence scores
            confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Extract text
            text = pytesseract.image_to_string(processed_image)
            
            return {
                "extracted_text": text.strip(),
                "confidence_score": round(avg_confidence, 1),
                "quality_level": get_quality_level(avg_confidence),
                "preprocessing_applied": True,
                "text_blocks": len(confidences),
                "processing_notes": get_processing_notes(avg_confidence),
                "file_type": "image"
            }
        
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = ""
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            return {
                "extracted_text": text.strip(),
                "confidence_score": 98.0,  # Word documents have very reliable text extraction
                "quality_level": "excellent",
                "preprocessing_applied": False,
                "text_blocks": paragraph_count,
                "processing_notes": "Word document text extraction - very high reliability",
                "file_type": "docx"
            }
        
        else:
            return {
                "extracted_text": "Unsupported file type for text extraction",
                "confidence_score": 0.0,
                "quality_level": "failed",
                "preprocessing_applied": False,
                "text_blocks": 0,
                "processing_notes": "File type not supported for text extraction",
                "file_type": "unsupported"
            }
            
    except Exception as e:
        logger.error(f"Error in enhanced OCR: {e}")
        return {
            "extracted_text": "OCR processing failed",
            "confidence_score": 0.0,
            "quality_level": "failed",
            "preprocessing_applied": False,
            "text_blocks": 0,
            "processing_notes": f"Error during text extraction: {str(e)}",
            "file_type": "error"
        }

def preprocess_image_for_ocr(image_bytes: bytes) -> Image:
    """Enhance image quality for better OCR results"""
    try:
        # Open image
        image = Image.open(io.BytesIO(image_bytes))
        
        # Convert to grayscale if not already
        if image.mode != 'L':
            image = image.convert('L')
        
        # Enhance contrast
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)  # Increase contrast by 50%
        
        # Reduce noise (simple approach)
        image = image.filter(ImageFilter.MedianFilter(size=3))
        
        # Sharpen text
        image = image.filter(ImageFilter.SHARPEN)
        
        return image
        
    except Exception as e:
        logger.error(f"Image preprocessing failed: {e}")
        # Return original image if preprocessing fails
        return Image.open(io.BytesIO(image_bytes))

def get_quality_level(confidence: float) -> str:
    """Determine quality level based on confidence score"""
    if confidence >= 90:
        return "excellent"
    elif confidence >= 70:
        return "good"
    elif confidence >= 50:
        return "fair"
    else:
        return "poor"

def get_processing_notes(confidence: float) -> str:
    """Generate user-friendly processing notes"""
    if confidence >= 90:
        return "High quality text extraction - ready for analysis"
    elif confidence >= 70:
        return "Good text extraction - minor quality issues detected"
    elif confidence >= 50:
        return "Fair text extraction - some text may need manual review"
    else:
        return "Poor text extraction - consider uploading a clearer image"

async def analyze_text_for_fraud(text: str) -> dict:
    """Analyze text for fraud patterns using AI models"""
    start_time = datetime.utcnow()
    
    try:
        fraud_score = 0.0
        patterns = []
        emotion_analysis = None
        pattern_analysis = None
        
        # 1. Emotion-based analysis using Hugging Face model
        if text_classifier and text_classifier != "limited":
            try:
                emotion_results = text_classifier(text)
                emotions = []
                fraud_indicators = []
                
                for result in emotion_results:
                    emotions.append({
                        "emotion": result["label"],
                        "confidence": result["score"]
                    })
                    
                    # Check for fraud-indicating emotions
                    if result["label"] in ["anger", "fear", "sadness"] and result["score"] > 0.3:
                        fraud_indicators.append({
                            "emotion": result["label"],
                            "confidence": result["score"],
                            "reason": f"Suspicious emotional tone detected: {result['label']}"
                        })
                        fraud_score += result["score"] * 0.2
                
                emotion_analysis = {
                    "emotions": emotions,
                    "fraud_indicators": fraud_indicators,
                    "emotion_fraud_score": sum([ind["confidence"] for ind in fraud_indicators]) / len(fraud_indicators) if fraud_indicators else 0.0,
                    "model_used": "cardiffnlp/twitter-roberta-base-emotion"
                }
            except Exception as e:
                logger.warning(f"Emotion analysis failed: {e}")
        
        # 2. Pattern-based fraud detection
        fraud_keywords = {
            "urgency": ["urgent", "immediate", "asap", "emergency", "critical", "rush"],
            "confidentiality": ["confidential", "secret", "do not share", "private", "exclusive"],
            "payment": ["wire transfer", "offshore", "bitcoin", "gift cards", "western union"],
            "amount": ["amount", "total", "sum", "cost", "price", "payment"]
        }
        
        text_lower = text.lower()
        pattern_scores = []
        
        for category, keywords in fraud_keywords.items():
            matches = [keyword for keyword in keywords if keyword in text_lower]
            if matches:
                score = len(matches) / len(keywords)
                pattern_scores.append({
                    "pattern": category,
                    "confidence": score,
                    "description": f"Detected {len(matches)} {category} indicators"
                })
                patterns.extend(matches)
                fraud_score += score * 0.3
        
        pattern_analysis = {
            "patterns": pattern_scores,
            "pattern_fraud_score": sum([p["confidence"] for p in pattern_scores]) / len(pattern_scores) if pattern_scores else 0.0
        }
        
        # 3. Text classification for additional fraud detection
        if text_classifier and text_classifier != "limited":
            try:
                classification_results = text_classifier(text)
                for result in classification_results:
                    if result["label"] in ["anger", "fear", "sadness"] and result["score"] > 0.5:
                        patterns.append(f"Suspicious emotional tone: {result['label']}")
                        pattern_scores.append({
                            "pattern": "emotional_manipulation",
                            "confidence": result["score"],
                            "description": f"Suspicious emotional tone detected: {result['label']}"
                        })
                        fraud_score += result["score"] * 0.2
            except Exception as e:
                logger.warning(f"Text classification failed: {e}")
        
        # Normalize fraud score to 0-1 range
        fraud_score = min(fraud_score, 1.0)
        
        # Determine risk level
        if fraud_score >= 0.7:
            risk_level = "HIGH"
        elif fraud_score >= 0.4:
            risk_level = "MEDIUM"
        else:
            risk_level = "LOW"
        
        processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
        
        return {
            "fraud_score": round(fraud_score, 3),
            "risk_level": risk_level,
            "patterns": list(set(patterns)),
            "emotion_analysis": emotion_analysis,
            "pattern_analysis": pattern_analysis,
            "processing_time": round(processing_time, 2)
        }
        
    except Exception as e:
        logger.error(f"Error in fraud analysis: {e}")
        return {
            "fraud_score": 0.0,
            "risk_level": "UNKNOWN",
            "patterns": [],
            "emotion_analysis": None,
            "pattern_analysis": None,
            "processing_time": 0.0
        }

if __name__ == "__main__":
    # Get server configuration
    server_config = config.get_server_config()
    
    # Log configuration
    logger.info(f"Starting FraudDocAI AI Service with configuration:")
    logger.info(f"  Host: {server_config['host']}")
    logger.info(f"  Port: {server_config['port']}")
    logger.info(f"  Reload: {server_config['reload']}")
    logger.info(f"  Log Level: {server_config['log_level']}")
    
    # Start the server
    uvicorn.run(
        "app:app",
        **server_config
    )
